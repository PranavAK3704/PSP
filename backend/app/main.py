"""Valmo Partner Support Platform — resolution engine API (BRD-aligned).

FastAPI service. Reactive resolution and proactive monitoring both stream their
pipeline trace over Server-Sent Events so the UI can watch the engine work live.

Auth: every /api route below is gated server-side (see the auth package). Two routes
are intentionally public — POST /api/auth/login and GET /api/health (Render health
check). Authoring writes need author-or-approver; approvals need approver.
"""
from __future__ import annotations

import csv
import io
import json
import os
from pathlib import Path

from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse

load_dotenv(Path(__file__).resolve().parents[1] / ".env")

from .audit import cpd                                 # noqa: E402
from .audit import rubric as audit_rubric              # noqa: E402
from .audit import runner as audit_runner             # noqa: E402
from .auth import store as auth_store                  # noqa: E402
from .auth import tokens as auth_tokens                # noqa: E402
from .auth.deps import current_user, require_role      # noqa: E402
from .channels import whatsapp                          # noqa: E402
from .engine import conversation, dispositions          # noqa: E402
from .knowledge import blueprints, governance, policies, sop_compiler, store  # noqa: E402
from .kt import engine as kt_engine                     # noqa: E402
from .l3 import platform as l3                          # noqa: E402
from .llm import registry as llm_registry               # noqa: E402
from .ledger import concern_log, trace_log               # noqa: E402
from .monitor import monitor                             # noqa: E402
from .substrate import captain_context as ctx            # noqa: E402
from .trust import constitution                          # noqa: E402

app = FastAPI(title="Valmo Partner Support Platform", version="demo-1.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# Role gates (server-side; the client is never trusted for role). An approver
# implicitly satisfies an author-level gate (approver ≥ author) — see auth/deps.py.
_authed = Depends(current_user)               # any authenticated user
_author = Depends(require_role("author"))      # authoring writes (author or approver)
_approver = Depends(require_role("approver"))  # approvals / go-live (approver only)


@app.on_event("startup")
def _seed_users():
    """First-run: seed the initial approver + baseline authored content (domain brains, the
    COD Shortfall SOP) so the library is populated on boot regardless of request order."""
    auth_store.seed_initial()
    try:
        from .kt import engine as kt_engine
        kt_engine.ensure_seeded()
        blueprints.load()   # seeds the Losses brain if the store is empty
    except Exception:  # noqa: BLE001 — never let seeding block startup
        pass


class ChatIn(BaseModel):
    captain_id: str
    message: str
    conversation_id: str | None = None
    channel: str = "chat"
    attachments: list[dict] | None = None   # [{filename, mime, size}] — thumbnail kept client-side


class CompileIn(BaseModel):
    sop_text: str


class SopApproveIn(BaseModel):
    policy: dict
    contributor: str = "sop-author"
    sop_id: str = ""            # set when editing an existing SOP → update in place (no duplicate)


class SopDeleteIn(BaseModel):
    sop_id: str


class BlueprintCompileIn(BaseModel):
    raw_text: str
    domain: str | None = None


class BlueprintSaveIn(BaseModel):
    blueprint: dict
    contributor: str = "author"


class BlueprintApproveIn(BaseModel):
    domain: str


class WhatsAppIn(BaseModel):
    from_: str | None = None
    text: str = ""
    payload: dict | None = None   # raw Meta webhook (optional)


class KtIn(BaseModel):
    text: str
    contributor: str = "anonymous"
    attachments: list[str] | None = None


class KtReviewIn(BaseModel):
    kt_id: str
    approve: bool
    reviewer: str = "reviewer"


class NuanceIn(BaseModel):
    text: str                              # plain-language rule / what to do
    domain: str = "general"                # payments | fe_id | losses_debits | cash_cod | consumables | orders | ...
    contributor: str = "sop-author"
    sop_ref: str = ""                      # the SOP/disposition this refines (optional)
    required_inputs: list[str] | None = None   # plain labels the captain must provide
    triggers: list[str] | None = None      # retrieval phrases (optional)
    from_concern_id: str = ""              # set when captured from a live case (L3 correction)
    resolves_gap_id: str = ""              # set when filling a captured auto_gap


class SatisfactionIn(BaseModel):
    concern_id: str
    captain_id: str
    satisfied: bool
    note: str = ""


class L3ResolveIn(BaseModel):
    concern_id: str
    resolution_note: str = ""
    resolver: str = "L3"


class AuditRubricIn(BaseModel):
    dimensions: list[dict]


class FrameworkSaveIn(BaseModel):
    framework: dict


class AuditRunIn(BaseModel):
    concern_id: str


class AuditBatchIn(BaseModel):
    limit: int = 10


class LoginIn(BaseModel):
    email: str
    password: str


class CreateUserIn(BaseModel):
    email: str
    name: str = ""
    role: str = "viewer"
    password: str


def _sse(gen):
    for event in gen:
        yield {"event": "trace", "data": json.dumps(event)}
    yield {"event": "end", "data": "{}"}


# ── Auth (login is PUBLIC; everything else needs a valid token) ──────────────
@app.post("/api/auth/login")
def login(body: LoginIn):
    """Public. Validate credentials → issue a 12h signed token + the user's public profile."""
    user = auth_store.verify_password(body.email, body.password)
    if not user:
        raise HTTPException(status_code=401, detail="invalid email or password")
    return {"token": auth_tokens.make_token(user["email"], user["role"]), "user": user}


@app.get("/api/auth/me")
def auth_me(user: dict = Depends(current_user)):
    """The current user (email, name, role) — resolved from the store, secrets stripped."""
    full = auth_store.get_user(user["email"])
    if not full:
        raise HTTPException(status_code=401, detail="user no longer exists")
    return {"email": full["email"], "name": full.get("name", ""), "role": full.get("role", "viewer")}


@app.get("/api/auth/users", dependencies=[_approver])
def auth_users():
    """Approver only. List users (email, name, role) — NEVER salts/hashes."""
    return {"users": auth_store.list_users()}


@app.post("/api/auth/users", dependencies=[_approver])
def auth_create_user(body: CreateUserIn):
    """Approver only. Create a team member with a temp password + role."""
    try:
        user = auth_store.create_user(body.email, body.name, body.role, body.password)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"ok": True, "user": user}


# ── Health (PUBLIC — Render health check) ────────────────────────────────────
@app.get("/api/health")
def health():
    from .substrate import loss_db
    ds = {"source": loss_db.source()}          # 'remote' (Turso) | 'local' (baked) | 'none'
    try:
        ds["losses"] = loss_db._query("SELECT COUNT(*) AS n FROM losses", ())[0]["n"]
    except Exception:  # noqa: BLE001
        ds["losses"] = None
    return {"ok": True, "provider": llm_registry.active_provider_name(),
            "knowledge": store.corpus_stats(), "data": ds}


@app.get("/api/captains", dependencies=[_authed])
def captains():
    return {"captains": ctx.known_captains()}


@app.get("/api/captain/{captain_id}", dependencies=[_authed])
def captain(captain_id: str):
    return ctx.get_context(captain_id) or {"error": "unknown captain"}


@app.post("/api/chat", dependencies=[_authed])
def chat(body: ChatIn):
    """Reactive resolution — stateful, multi-turn. Streams the trace as SSE.

    Pass a stable conversation_id across turns; the engine may end a turn in a
    `need_input` event (asking the partner) and resume on the next turn.
    """
    import uuid as _uuid
    conv_id = body.conversation_id or ("conv-" + _uuid.uuid4().hex[:10])
    return EventSourceResponse(_sse(
        conversation.handle_turn(conv_id, body.captain_id, body.message, body.channel,
                                 attachments=body.attachments)))


@app.get("/api/monitor/{captain_id}", dependencies=[_authed])
def monitor_scan(captain_id: str):
    """Proactive monitoring — streams detect->nudge trace as SSE."""
    return EventSourceResponse(_sse(monitor.scan_captain(captain_id)))


@app.get("/api/dispositions", dependencies=[_authed])
def dispositions_list():
    return {"dispositions": dispositions.catalogue()}


@app.get("/api/policies", dependencies=[_authed])
def policies_list():
    return {"policies": policies.all_policies()}


@app.post("/api/sop/compile", dependencies=[_author])
def compile_sop(body: CompileIn):
    """Streams the compilation as SSE stages so the UI can animate the real structuring/tiering.
    The final `stage: done` event carries the full ExecutablePolicy + gaps (inline) + meta."""
    return EventSourceResponse(_sse(sop_compiler.compile_sop_streamed(body.sop_text)))


@app.post("/api/sop/extract", dependencies=[_author])
async def sop_extract(file: UploadFile = File(...)):
    """Upload a real ops artifact (Excel / Word / PDF / CSV / text) → extract its content as
    readable text (tables flattened) → return {text, source_name}. The author reviews/edits the
    text, then Compiles it through the normal SOP pipeline. The machine structures; a human
    approves. Extraction only — no LLM, no state write."""
    raw = await file.read()
    text = _extract_text(raw, file.content_type, file.filename)
    return {"text": text, "source_name": file.filename}


@app.post("/api/sop/save", dependencies=[_author])
def save_sop(body: SopApproveIn):
    """Save a compiled SOP as a DRAFT (author-or-approver) so it is never lost — it shows in
    the library and can be approved later. sop_id set → edit in place. Returns {ok, id, gaps}."""
    entry = sop_compiler.save_sop_draft(body.policy, body.contributor, body.sop_id)
    return {"ok": True, "id": entry["id"], "gaps": sop_compiler.detect_policy_gaps(body.policy)}


@app.post("/api/sop/approve", dependencies=[_approver])
def approve_sop(body: SopApproveIn):
    """A reviewed structured SOP enters the retrieval corpus (with reload) so the engine
    follows it. sop_id set → update the existing SOP in place. Returns {ok, id, gaps}."""
    entry = sop_compiler.approve_sop(body.policy, body.contributor, body.sop_id)
    return {"ok": True, "id": entry["id"], "gaps": sop_compiler.detect_policy_gaps(body.policy)}


@app.post("/api/sop/delete", dependencies=[_approver])
def delete_sop(body: SopDeleteIn):
    """Remove a compiled SOP by id (Knowledge Base management, approver only). Returns {ok, removed}."""
    removed = sop_compiler.delete_sop(body.sop_id)
    return {"ok": True, "removed": removed}


# ── Authoring Studio: Domain Blueprints (a domain's stage-0 "brain") ─────────
@app.post("/api/blueprint/compile", dependencies=[_author])
def compile_blueprint(body: BlueprintCompileIn):
    """Streams the structuring of a free-text domain walkthrough into a Blueprint (SSE stages).
    The final `stage: done` event carries {blueprint, gaps} — gaps surfaced AT creation (req #4)."""
    return EventSourceResponse(_sse(
        blueprints.compile_blueprint_streamed(body.raw_text, body.domain)))


@app.get("/api/blueprints", dependencies=[_authed])
def blueprints_list():
    """All Blueprints (draft + approved), newest-updated first, each with its inline gaps."""
    items = blueprints.list_blueprints()
    return {"blueprints": [{**b, "gaps": blueprints.detect_gaps(b)} for b in items]}


@app.post("/api/blueprint/save", dependencies=[_author])
def save_blueprint(body: BlueprintSaveIn):
    """Save/replace a Blueprint as a draft. Returns {ok, blueprint, gaps, existing_brain}.
    existing_brain is the dedup summary of any blueprint already stored for this domain —
    resolved BEFORE the save (save overwrites the one-per-domain entry)."""
    domain = str((body.blueprint or {}).get("domain", "")).strip().lower()
    existing = blueprints.existing_brain(domain)
    bp = blueprints.save(body.blueprint, body.contributor)
    return {"ok": True, "blueprint": bp, "gaps": blueprints.detect_gaps(bp),
            "existing_brain": existing}


@app.post("/api/blueprint/approve", dependencies=[_approver])
def approve_blueprint(body: BlueprintApproveIn):
    """Approve a Blueprint → status=approved + reload; the engine now follows it. Returns {ok}."""
    bp = blueprints.approve(body.domain)
    if bp is None:
        return {"ok": False, "error": "blueprint not found"}
    return {"ok": True, "domain": bp["domain"], "status": bp["status"]}


@app.get("/api/ledger", dependencies=[_authed])
def ledger():
    return {"concerns": concern_log.all_concerns(), "stats": concern_log.stats()}


@app.get("/api/concern/{concern_id}/trace", dependencies=[_authed])
def concern_trace(concern_id: str):
    """The persisted resolution TRACE for a concern — the stage-by-stage events
    (node/label/status/detail/data) the engine streamed while resolving it.
    Returns {events: []} if no trace was stored (e.g. pre-dating the feature)."""
    rec = trace_log.get(concern_id)
    if not rec:
        return {"concern_id": concern_id, "events": []}
    return rec


# CSV column order for the flat ledger export.
_EXPORT_COLUMNS = ["id", "seq", "logged_at", "captain_id", "conversation_id",
                   "disposition", "action_taken", "outcome", "intent", "amount_inr",
                   "escalation_team", "reply"]


@app.get("/api/ledger/export", dependencies=[_authed])
def ledger_export(format: str = "json"):
    """Download the Concern Log. format=csv → flat columns; format=json → full records.
    Sets Content-Disposition so the browser downloads it as a file."""
    concerns = concern_log.all_concerns()
    fmt = (format or "json").lower()
    if fmt == "csv":
        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=_EXPORT_COLUMNS, extrasaction="ignore")
        writer.writeheader()
        for c in concerns:
            row = {k: c.get(k, "") for k in _EXPORT_COLUMNS}
            # flatten any newlines in free-text so the CSV stays one-row-per-concern
            if row.get("reply"):
                row["reply"] = str(row["reply"]).replace("\n", " ").replace("\r", " ")
            if row.get("intent"):
                row["intent"] = str(row["intent"]).replace("\n", " ").replace("\r", " ")
            writer.writerow(row)
        return Response(
            content=buf.getvalue(),
            media_type="text/csv",
            headers={"Content-Disposition": 'attachment; filename="concern_log.csv"'},
        )
    return Response(
        content=json.dumps(concerns, indent=1),
        media_type="application/json",
        headers={"Content-Disposition": 'attachment; filename="concern_log.json"'},
    )


@app.get("/api/constitution", dependencies=[_authed])
def get_constitution():
    return {"principles": constitution.PRINCIPLES, "tiers": constitution.TIERS}


# ── Channels: WhatsApp (the near-term primary channel) ──────────────────────
@app.post("/api/whatsapp/webhook", dependencies=[_authed])
def whatsapp_webhook(body: WhatsAppIn):
    """WhatsApp version of the engine. Parses a webhook, runs a conversation turn
    keyed by phone (so multi-turn works over WhatsApp), and returns the reply."""
    payload = body.payload or {"from": body.from_, "text": body.text}
    parsed = whatsapp.parse_webhook(payload)
    if not parsed or not parsed.get("captain_id"):
        return {"error": "unknown sender or empty message", "hint": "map phone→captain in whatsapp.PHONE_TO_CAPTAIN"}
    conv_id = f"wa-{parsed['phone']}"
    reply, terminal = None, None
    # handle_turn ends a turn with a `reply` node carrying the captain-facing message
    # (same shape as the SSE 'reply' event): ev['data']['reply'], with ev['detail'] as
    # the plain-text fallback. Capture that so we forward the engine's ACTUAL reply.
    for ev in conversation.handle_turn(conv_id, parsed["captain_id"], parsed["text"], channel="whatsapp"):
        if ev["node"] == "reply":
            terminal = ev
            reply = ev["data"].get("reply") or ev.get("detail")
    # whatsapp.send is a STUB — pending real WhatsApp Business API wiring (see channels/whatsapp.py).
    outbound = whatsapp.send(parsed["phone"], reply or "…")
    return {"captain_id": parsed["captain_id"], "state": terminal["node"] if terminal else "?",
            "reply": reply, "outbound": outbound}


# ── L3 functional-team platform ─────────────────────────────────────────────
@app.get("/api/l3/inbox", dependencies=[_authed])
def l3_inbox():
    return {"items": l3.inbox(), "teams": l3.team_metrics()}


@app.post("/api/l3/resolve", dependencies=[_authed])
def l3_resolve(body: L3ResolveIn):
    """L3 resolves an escalated case → drops it from the active queue and creates a
    captain-facing follow-up (closes the loop across states)."""
    return l3.resolve(body.concern_id, body.resolution_note, body.resolver)


@app.get("/api/captain/{captain_id}/cases", dependencies=[_authed])
def captain_cases(captain_id: str):
    """Captain-facing 'My Cases': escalated cases + live status + resolution (polled by the widget)."""
    return {"cases": l3.cases(captain_id)}


# ── KT engine ───────────────────────────────────────────────────────────────
@app.post("/api/kt/submit", dependencies=[_author])
def kt_submit(body: KtIn):
    return kt_engine.submit(body.text, body.contributor, body.attachments)


@app.get("/api/kt", dependencies=[_authed])
def kt_list():
    return {"pending": kt_engine.pending(), "all": kt_engine.all_kt()}


@app.post("/api/kt/review", dependencies=[_approver])
def kt_review(body: KtReviewIn):
    """Approver only — approving/rejecting KT changes what the engine follows (go-live)."""
    res = kt_engine.review(body.kt_id, body.approve, body.reviewer)
    return res or {"error": "kt not found"}


# ── SOP required-data gaps + non-tech nuance authoring ──────────────────────
@app.get("/api/sop/gaps", dependencies=[_authed])
def sop_gaps():
    """SOPs the engine can't fully act on because it doesn't know what to ask the captain,
    plus auto-captured gaps from escalations. Powers the Support Console authoring panel."""
    from .knowledge import gaps
    return gaps.detect()


@app.post("/api/sop/nuance", dependencies=[_author])
def sop_nuance(body: NuanceIn):
    """A non-tech author fills a gap / adds a correction. Enters the KT approval queue;
    once approved it becomes engine behaviour via retrieval — no code change."""
    return kt_engine.submit_nuance(
        body.text, body.domain, body.contributor, sop_ref=body.sop_ref,
        required_inputs=body.required_inputs, triggers=body.triggers,
        from_concern_id=body.from_concern_id, resolves_gap_id=body.resolves_gap_id)


# ── Audit + CPD + satisfaction ──────────────────────────────────────────────
@app.post("/api/satisfaction", dependencies=[_authed])
def satisfaction(body: SatisfactionIn):
    return cpd.record_satisfaction(body.concern_id, body.captain_id, body.satisfied, body.note)


@app.get("/api/audit", dependencies=[_authed])
def audit():
    return {"trail": cpd.audit_trail(), "cpd": cpd.cpd_items(), "satisfaction": cpd.satisfaction_stats()}


# ── Auditing Studio: editable rubric + LLM judge + score dashboard ───────────
@app.get("/api/audit/rubric", dependencies=[_authed])
def audit_rubric_get():
    """The current (versioned) audit rubric the LLM judge scores against."""
    return audit_rubric.get_rubric()


@app.post("/api/audit/rubric", dependencies=[_author])
def audit_rubric_save(body: AuditRubricIn):
    """Save an edited rubric → bumps version. Returns the new rubric. Authoring write."""
    return audit_rubric.save_rubric(body.dimensions)


@app.post("/api/audit/run", dependencies=[_authed])
def audit_run(body: AuditRunIn):
    """Audit one concern with the LLM judge → composite + per-dimension scores."""
    return audit_runner.audit_concern(body.concern_id)


@app.post("/api/audit/run_batch", dependencies=[_authed])
def audit_run_batch(body: AuditBatchIn):
    """Audit the most recent N un-audited concerns (sampling). Returns a summary."""
    return audit_runner.audit_batch(body.limit)


@app.get("/api/audit/scores", dependencies=[_authed])
def audit_scores():
    """Audit history + aggregates (avg composite, per-dimension avg, trend, by disposition)."""
    return audit_runner.scores()


# ── Auditing Studio: dynamic, editable Governance Framework ──────────────────
def _extract_text(raw: bytes, content_type: str, filename: str) -> str:
    """Best-effort text extraction from an uploaded doc, so an author can drop a real
    operations artifact (SOP sheet, KT doc, framework) straight in and let the machine
    structure it. Handles: PDF (pypdf), Excel .xlsx (openpyxl), Word .docx (python-docx),
    and utf-8 text/csv/markdown/plain. Tables are flattened to readable rows so the LLM
    keeps the structure. Raises 400 on empty/unreadable content."""
    if raw and len(raw) > 12_000_000:   # ~12MB cap — a huge upload read fully into memory could OOM
        raise HTTPException(status_code=413, detail="File too large — please upload under ~12MB.")
    import io as _io
    ctype = (content_type or "").lower()
    name = (filename or "").lower()

    def _fail(msg):
        raise HTTPException(status_code=400, detail=msg)

    if "application/pdf" in ctype or name.endswith(".pdf"):
        from pypdf import PdfReader
        try:
            reader = PdfReader(_io.BytesIO(raw))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as e:  # noqa: BLE001
            _fail(f"could not read PDF: {e}")
    elif name.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in ctype:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(_io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"## Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        except Exception as e:  # noqa: BLE001
            _fail(f"could not read Excel file: {e}")
    elif name.endswith((".docx",)) or "wordprocessingml" in ctype:
        try:
            import docx  # python-docx
            doc = docx.Document(_io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        except Exception as e:  # noqa: BLE001
            _fail(f"could not read Word file: {e}")
    elif name.endswith((".xls", ".doc")):
        _fail("legacy .xls/.doc isn't supported — please re-save as .xlsx / .docx (or paste the text).")
    else:
        text = raw.decode("utf-8", errors="replace")

    text = (text or "").strip()
    if not text:
        _fail("no extractable text in the uploaded file")
    return text


@app.get("/api/framework", dependencies=[_authed])
def framework_get():
    """The current (versioned) Governance Framework — the seeded PLACEHOLDER on first run."""
    return governance.get()


@app.post("/api/framework", dependencies=[_author])
def framework_save(body: FrameworkSaveIn):
    """Save an edited framework as a draft (author-or-approver). Returns {ok, framework}."""
    fw = governance.save(body.framework)
    return {"ok": True, "framework": fw}


@app.post("/api/framework/upload", dependencies=[_author])
async def framework_upload(file: UploadFile = File(...)):
    """Upload a framework document → extract text (pypdf for PDF; utf-8 for text/csv/md/plain)
    → structure it into the Framework model with the LLM → return {framework(draft), source_name}.
    The machine structures it; a human reviews before approving (author-or-approver)."""
    raw = await file.read()
    text = _extract_text(raw, file.content_type, file.filename)
    fw = governance.structure_framework_from_text(text)
    return {"framework": fw, "source_name": file.filename}


@app.post("/api/framework/approve", dependencies=[_approver])
def framework_approve():
    """Publish the current framework → status=approved + bump version (approver only)."""
    fw = governance.approve()
    return {"ok": True, "framework": fw}


@app.get("/api/insights", dependencies=[_authed])
def insights():
    """One aggregate for the ops control tower."""
    # Additive ops metrics — each guarded so a hiccup here can never blank the panel.
    try:
        active_breaches = sum(1 for it in l3.inbox() if it.get("breached"))
    except Exception:  # noqa: BLE001
        active_breaches = 0
    try:
        avg_resolution_time = concern_log.resolution_time_stats()
    except Exception:  # noqa: BLE001
        avg_resolution_time = {"display": "—", "hours": None, "sample": 0, "basis": "n/a"}
    return {
        "provider": llm_registry.active_provider_name(),
        "knowledge": store.corpus_stats(),
        "ledger": concern_log.stats(),
        "satisfaction": cpd.satisfaction_stats(),
        "l3_teams": l3.team_metrics(),
        "cpd_open": len(cpd.cpd_items()),
        "dispositions": len(dispositions.catalogue()),
        "active_breaches": active_breaches,
        "avg_resolution_time": avg_resolution_time,
    }


# ── Serve the built frontend (production / Render) ───────────────────────────
# The Vite build lands in <repo>/frontend/dist (local) or /app/frontend/dist (Docker).
# We mount it AFTER every /api route so the API always wins; a missing dist (local dev
# via `run.sh`, where Vite serves the UI on :5190) simply skips mounting so boot never
# breaks. Unknown non-/api GET paths fall back to index.html (SPA-friendly).
_DIST = Path(os.environ.get("PSP_STATIC_DIR") or (Path(__file__).resolve().parents[2] / "frontend" / "dist"))
if (_DIST / "index.html").is_file():
    _assets = _DIST / "assets"
    if _assets.is_dir():
        app.mount("/assets", StaticFiles(directory=str(_assets)), name="assets")

    @app.get("/", include_in_schema=False)
    def _spa_root():
        return FileResponse(str(_DIST / "index.html"))

    @app.get("/{full_path:path}", include_in_schema=False)
    def _spa_fallback(full_path: str):
        # never let the SPA catch-all shadow the API surface
        if full_path == "api" or full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="not found")
        candidate = _DIST / full_path
        if candidate.is_file():
            return FileResponse(str(candidate))
        return FileResponse(str(_DIST / "index.html"))
